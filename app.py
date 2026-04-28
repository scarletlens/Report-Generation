from __future__ import annotations

import copy
import csv
import ipaddress
import io
import json
import os
import re
from collections import OrderedDict
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, quote, urlparse
from wsgiref.simple_server import WSGIRequestHandler, make_server

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


BASE_DIR = Path(__file__).resolve().parent
MATERIAL_DIR = BASE_DIR / "material"
DATA_FILE = MATERIAL_DIR / "报表数据.xlsx"
TEMPLATE_FILE = MATERIAL_DIR / "报表模板.docx"
RESOURCE_FILE = MATERIAL_DIR / "资源清单.xls"
EXTRA_RESOURCE_DIR = MATERIAL_DIR / "其他资源清单"
EXTRA_UTILIZATION_DIR = MATERIAL_DIR / "其他利用率"

SHEET_NAME = "利用率"
SECURITY_SHEET_NAME = "安全态势"
ADDRESS_SHEET_NAME = "地址段"
HEADERS = ["厅局名称", "系统名称", "IP", "时间", "CPU使用率", "内存使用率", "磁盘使用率"]
REGION_XC = "信创"
REGION_NON_XC = "非信创"
REGIONS = [REGION_XC, REGION_NON_XC]

FONT_EAST_ASIA = "微软雅黑"
FONT_LATIN = "Aptos"
COLOR_INK = "1F2937"
COLOR_MUTED = "667085"
COLOR_ACCENT = "1F6F68"
COLOR_ACCENT_DARK = "174F4B"
COLOR_HEADER_BG = "1F4E5F"
COLOR_HEADER_TEXT = "FFFFFF"
COLOR_ROW_ALT = "F3F8F7"
COLOR_BORDER = "D7DEE8"
COLOR_PAGE_BG = "F7FAF9"
COLOR_WAVE = "DCEDEA"
COLOR_WAVE_LIGHT = "EEF7F5"

_CACHE: dict[str, Any] = {
    "mtime": None,
    "rows": None,
    "csv_rows": None,
    "security_rows": None,
    "resource_rows": None,
    "resource_index": None,
    "address_segments": None,
}


def iter_block_items(parent: Document):
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def normalize_date(value: Any) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y/%m/%d %H:%M:%S", "%Y-%m-%d", "%Y/%m/%d"):
        try:
            return datetime.strptime(text, fmt).strftime("%Y-%m-%d")
        except ValueError:
            pass
    return text


def normalize_number(value: Any) -> str:
    if value in (None, ""):
        return "-"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value).strip() or "-"
    if number.is_integer():
        return str(int(number))
    return f"{number:.2f}"


def sum_numbers(*values: Any) -> float | None:
    total = 0.0
    has_value = False
    for value in values:
        if value in (None, ""):
            continue
        try:
            number = float(value)
        except (TypeError, ValueError):
            continue
        total += number
        has_value = True
    return total if has_value else None


def display_month(value: str) -> str:
    try:
        dt = datetime.strptime(value, "%Y-%m-%d")
        return f"{dt.year}年{dt.month:02d}月"
    except ValueError:
        return value


def parse_date(value: str) -> date | None:
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return None


def month_bounds(month: str) -> tuple[str, str]:
    try:
        start = datetime.strptime(month, "%Y-%m").date()
    except (TypeError, ValueError):
        fallback = datetime.today().date().replace(day=1)
        start = fallback
    if start.month == 12:
        next_month = start.replace(year=start.year + 1, month=1, day=1)
    else:
        next_month = start.replace(month=start.month + 1, day=1)
    end = next_month - timedelta(days=1)
    return start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d")


def normalize_period(
    mode: str | None = None,
    month: str | None = None,
    start: str | None = None,
    end: str | None = None,
    date_value: str | None = None,
) -> dict[str, str]:
    if mode == "range":
        start_value = normalize_date(start or date_value)
        end_value = normalize_date(end or start_value)
    else:
        month_value = month or (normalize_date(date_value)[:7] if date_value else "")
        start_value, end_value = month_bounds(month_value)
        mode = "month"

    if parse_date(start_value) and parse_date(end_value) and parse_date(start_value) > parse_date(end_value):
        start_value, end_value = end_value, start_value

    return {
        "mode": mode or "month",
        "start": start_value,
        "end": end_value,
        "month": start_value[:7],
        "label": period_label(start_value, end_value),
        "filename": period_filename(start_value, end_value),
    }


def period_label(start: str, end: str) -> str:
    if start == end:
        return display_month(start) if start.endswith("-01") else start
    if start[:7] == end[:7]:
        month_start, month_end = month_bounds(start[:7])
        if start == month_start and end == month_end:
            return display_month(start)
    return f"{start}至{end}"


def period_filename(start: str, end: str) -> str:
    if start == end:
        return start
    if start[:7] == end[:7]:
        month_start, month_end = month_bounds(start[:7])
        if start == month_start and end == month_end:
            return display_month(start)
    return f"{start}至{end}"


def date_in_period(value: str, period: dict[str, str]) -> bool:
    current = parse_date(value)
    start = parse_date(period["start"])
    end = parse_date(period["end"])
    if not current or not start or not end:
        return value == period["start"] or value == period["end"]
    return start <= current <= end


def metric(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, (int, float)):
        return f"{value:.2f}"
    text = str(value).strip()
    if not text:
        return "-"
    try:
        return f"{float(text):.2f}"
    except ValueError:
        return text


def normalize_header(value: Any) -> str:
    return str(value).strip() if value is not None else ""


def first_present(record: dict[str, Any], names: list[str]) -> Any:
    for name in names:
        if name in record and record[name] not in (None, ""):
            return record[name]
    return ""


def extract_ip(value: Any) -> str:
    text = "" if value is None else str(value)
    match = re.search(r"(?:\d{1,3}\.){3}\d{1,3}", text.replace(" ", ""))
    return match.group(0) if match else ""


def parse_network(value: Any) -> ipaddress.IPv4Network | None:
    text = "" if value is None else str(value)
    text = re.sub(r"\s+", "", text)
    if not text:
        return None
    try:
        return ipaddress.ip_network(text, strict=False)
    except ValueError:
        return None


def in_same_month(value: str, report_date: str) -> bool:
    try:
        return datetime.strptime(value, "%Y-%m-%d").strftime("%Y-%m") == datetime.strptime(report_date, "%Y-%m-%d").strftime("%Y-%m")
    except ValueError:
        return value == report_date


def normalize_regions(regions: list[str] | None) -> list[str]:
    selected = [region for region in (regions or REGIONS) if region in REGIONS]
    return selected or REGIONS


def region_label(regions: list[str] | None) -> str:
    selected = normalize_regions(regions)
    if set(selected) == set(REGIONS):
        return "全区域"
    return "、".join(selected)


def bureau_matches(value: str, selected: str) -> bool:
    value = normalize_header(value)
    selected = normalize_header(selected)
    return bool(value and selected and (value == selected or selected in value or value in selected))


def is_supported_excel(path: Path) -> bool:
    return path.suffix.lower() in (".xlsx", ".xlsm", ".xls")


def is_ignored_data_file(path: Path) -> bool:
    return path.name.startswith(".") or path.name.startswith("~") or path.name.startswith(".~")


def normalize_security_severity(severity: str, attack_result: str) -> str:
    severity = normalize_header(severity) or "-"
    attack_result = normalize_header(attack_result)
    success_words = ("成功", "失陷", "沦陷", "已攻陷", "入侵")
    unsuccessful_words = ("尝试", "失败", "未成功", "阻断", "拦截", "未遂")

    if any(word in attack_result for word in success_words):
        return severity
    if not any(word in attack_result for word in unsuccessful_words):
        return severity

    if severity in ("严重", "高危", "高"):
        return "中危"
    if severity in ("中危", "中"):
        return "低危"
    return severity if severity != "-" else "低危"


def cache_mtime() -> tuple[
    float,
    float | None,
    tuple[tuple[str, float], ...],
    tuple[tuple[str, float], ...],
    tuple[tuple[str, float], ...],
]:
    resource_mtime = RESOURCE_FILE.stat().st_mtime if RESOURCE_FILE.exists() else None
    utilization_mtimes = tuple(
        sorted(
            (str(path.relative_to(MATERIAL_DIR)), path.stat().st_mtime)
            for path in iter_utilization_files()
        )
    )
    extra_resource_mtimes = tuple(
        sorted(
            (str(path.relative_to(MATERIAL_DIR)), path.stat().st_mtime)
            for path in EXTRA_RESOURCE_DIR.rglob("*")
            if path.is_file() and not path.name.startswith(".") and not path.name.startswith("~")
        )
    )
    return DATA_FILE.stat().st_mtime, resource_mtime, utilization_mtimes, extra_resource_mtimes


def reset_cache_if_needed() -> None:
    mtime = cache_mtime()
    if _CACHE["mtime"] != mtime:
        _CACHE.update(
            {
                "mtime": mtime,
                "rows": None,
                "csv_rows": None,
                "security_rows": None,
                "resource_rows": None,
                "resource_index": None,
                "address_segments": None,
            }
        )


def load_address_segments_uncached() -> list[dict[str, Any]]:
    workbook = load_workbook(DATA_FILE, data_only=True)
    if ADDRESS_SHEET_NAME not in workbook.sheetnames:
        return []

    sheet = workbook[ADDRESS_SHEET_NAME]
    headers = [normalize_header(cell.value) for cell in sheet[1]]
    category_names = ["区域", "类型", "类别", "信创类型", "是否信创", "区域类型"]
    rows: list[dict[str, Any]] = []

    for values in sheet.iter_rows(min_row=2, values_only=True):
        record = dict(zip(headers, values))
        category = normalize_header(first_present(record, category_names))
        explicit_region = ""
        if "非" in category and "信创" in category:
            explicit_region = REGION_NON_XC
        elif "信创" in category:
            explicit_region = REGION_XC

        for name, value in record.items():
            if "地址段" not in name and name not in ("网段", "IP网段"):
                continue
            network = parse_network(value)
            if not network:
                continue
            if explicit_region:
                region = explicit_region
            elif name == "政务外网地址段":
                region = REGION_XC
            else:
                region = REGION_NON_XC
            rows.append({"network": network, "region": region})
    return rows


def load_address_segments() -> list[dict[str, Any]]:
    reset_cache_if_needed()
    if _CACHE["address_segments"] is None:
        _CACHE["address_segments"] = load_address_segments_uncached()
    return _CACHE["address_segments"]


def classify_ip_region(ip: str, segments: list[dict[str, Any]] | None = None) -> str:
    if not ip:
        return REGION_NON_XC
    try:
        address = ipaddress.ip_address(ip)
    except ValueError:
        return REGION_NON_XC
    for segment in segments if segments is not None else load_address_segments():
        if address in segment["network"]:
            return segment["region"]
    return REGION_NON_XC


def row_region(row: dict[str, Any], segments: list[dict[str, Any]] | None = None) -> str:
    return classify_ip_region(row.get("ip", ""), segments)


def load_rows_uncached() -> list[dict[str, Any]]:
    workbook = load_workbook(DATA_FILE, data_only=True)
    if SHEET_NAME not in workbook.sheetnames:
        return []
    sheet = workbook[SHEET_NAME]
    headers = [normalize_header(cell.value) for cell in sheet[1]]

    missing = [name for name in HEADERS if name not in headers]
    if missing:
        raise ValueError(f"Excel 缺少必要列：{', '.join(missing)}")

    rows: list[dict[str, Any]] = []
    for values in sheet.iter_rows(min_row=2, values_only=True):
        record = dict(zip(headers, values))
        if not any(record.get(name) for name in HEADERS):
            continue
        rows.append(
            {
                "bureau": str(record["厅局名称"]).strip(),
                "system": str(record["系统名称"]).strip(),
                "ip": str(record["IP"]).strip(),
                "date": normalize_date(record["时间"]),
                "cpu": record["CPU使用率"],
                "memory": record["内存使用率"],
                "disk": record["磁盘使用率"],
            }
        )
    return rows


def load_rows() -> list[dict[str, Any]]:
    reset_cache_if_needed()
    if _CACHE["rows"] is None:
        _CACHE["rows"] = load_rows_uncached()
    return _CACHE["rows"]


def utilization_file_month(path: Path) -> str | None:
    stem = normalize_header(path.stem)
    match = re.search(r"(\d{4})\s*年\s*(\d{1,2})\s*月", stem)
    if not match:
        return None
    return f"{int(match.group(1)):04d}-{int(match.group(2)):02d}"


def month_from_csv_name(path: Path) -> str | None:
    return utilization_file_month(path)


def iter_utilization_files() -> list[Path]:
    paths: list[Path] = []
    for path in MATERIAL_DIR.iterdir():
        if not path.is_file() or is_ignored_data_file(path):
            continue
        if not utilization_file_month(path):
            continue
        if path.suffix.lower() == ".csv" or is_supported_excel(path):
            paths.append(path)
    if EXTRA_UTILIZATION_DIR.exists():
        for path in EXTRA_UTILIZATION_DIR.rglob("*"):
            if not path.is_file() or is_ignored_data_file(path):
                continue
            if utilization_file_month(path) and (path.suffix.lower() == ".csv" or is_supported_excel(path)):
                paths.append(path)
    return sorted(paths)


def utilization_row_from_values(
    values: list[Any],
    month: str,
    resource_index: dict[str, dict[str, str]],
) -> dict[str, Any] | None:
    if len(values) < 5:
        return None
    ip = extract_ip(values[1])
    if not ip:
        return None
    owner = resource_index.get(ip)
    if not owner:
        return None
    return {
        "bureau": owner["bureau"],
        "system": owner["system"],
        "ip": ip,
        "date": f"{month}-01",
        "month": month,
        "cpu": values[2],
        "memory": values[3],
        "disk": values[4],
    }


def load_csv_utilization_file(
    path: Path,
    month: str,
    resource_index: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.reader(file)
        for record in reader:
            row = utilization_row_from_values(record, month, resource_index)
            if row:
                rows.append(row)
    return rows


def load_xlsx_utilization_file(
    path: Path,
    month: str,
    resource_index: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    rows: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        for values in sheet.iter_rows(min_row=1, values_only=True):
            row = utilization_row_from_values(list(values), month, resource_index)
            if row:
                rows.append(row)
    return rows


def load_xls_utilization_file(
    path: Path,
    month: str,
    resource_index: dict[str, dict[str, str]],
) -> list[dict[str, Any]]:
    try:
        import xlrd
    except ImportError as exc:
        raise RuntimeError("缺少读取 .xls 利用率文件的依赖，请安装：pip install xlrd") from exc
    workbook = xlrd.open_workbook(str(path))
    rows: list[dict[str, Any]] = []
    for sheet in workbook.sheets():
        for row_index in range(sheet.nrows):
            values = [sheet.cell_value(row_index, col) for col in range(sheet.ncols)]
            row = utilization_row_from_values(values, month, resource_index)
            if row:
                rows.append(row)
    return rows


def load_csv_utilization_rows_uncached() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    resource_index = resource_index_by_ip()
    for path in iter_utilization_files():
        month = utilization_file_month(path)
        if not month:
            continue
        try:
            if path.suffix.lower() == ".csv":
                rows.extend(load_csv_utilization_file(path, month, resource_index))
            elif path.suffix.lower() in (".xlsx", ".xlsm"):
                rows.extend(load_xlsx_utilization_file(path, month, resource_index))
            elif path.suffix.lower() == ".xls":
                rows.extend(load_xls_utilization_file(path, month, resource_index))
        except Exception as exc:
            print(f"跳过利用率文件 {path.name}：{exc}")
    return rows


def load_csv_utilization_rows() -> list[dict[str, Any]]:
    reset_cache_if_needed()
    if _CACHE["csv_rows"] is None:
        _CACHE["csv_rows"] = load_csv_utilization_rows_uncached()
    return _CACHE["csv_rows"]


def aggregate_utilization_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: OrderedDict[tuple[str, str, str, str], dict[str, Any]] = OrderedDict()
    for row in rows:
        key = (row["bureau"], row["system"], row["ip"], row["date"], row["region"])
        grouped[key] = dict(row)
    if len({row["date"] for row in rows}) <= 1:
        return list(grouped.values())

    summary: OrderedDict[tuple[str, str, str, str], dict[str, Any]] = OrderedDict()
    for row in rows:
        key = (row["bureau"], row["system"], row["ip"], row["region"])
        item = summary.setdefault(
            key,
            {
                "bureau": row["bureau"],
                "system": row["system"],
                "ip": row["ip"],
                "region": row["region"],
                "date": "",
                "cpu_values": [],
                "memory_values": [],
                "disk_values": [],
            },
        )
        for metric_name, values_name in (("cpu", "cpu_values"), ("memory", "memory_values"), ("disk", "disk_values")):
            try:
                item[values_name].append(float(row[metric_name]))
            except (TypeError, ValueError):
                pass

    result = []
    for item in summary.values():
        result.append(
            {
                "bureau": item["bureau"],
                "system": item["system"],
                "ip": item["ip"],
                "region": item["region"],
                "date": "",
                "cpu": sum(item["cpu_values"]) / len(item["cpu_values"]) if item["cpu_values"] else None,
                "memory": sum(item["memory_values"]) / len(item["memory_values"]) if item["memory_values"] else None,
                "disk": sum(item["disk_values"]) / len(item["disk_values"]) if item["disk_values"] else None,
            }
        )
    return result


def months_in_period(period: dict[str, str]) -> list[str]:
    start = parse_date(period["start"])
    end = parse_date(period["end"])
    if not start or not end:
        return [period["month"]]
    current = start.replace(day=1)
    months = []
    while current <= end:
        months.append(current.strftime("%Y-%m"))
        if current.month == 12:
            current = current.replace(year=current.year + 1, month=1)
        else:
            current = current.replace(month=current.month + 1)
    return months


def filtered_rows(bureau: str, period: dict[str, str], regions: list[str] | None = None) -> list[dict[str, Any]]:
    selected_regions = normalize_regions(regions)
    segments = load_address_segments()
    rows = []
    csv_rows = load_csv_utilization_rows()
    excel_rows = load_rows()
    for month in months_in_period(period):
        month_period = normalize_period(mode="month", month=month)
        monthly_csv = [row for row in csv_rows if bureau_matches(row["bureau"], bureau) and row["month"] == month]
        source_rows = monthly_csv if monthly_csv else [
            row
            for row in excel_rows
            if bureau_matches(row["bureau"], bureau) and date_in_period(row["date"], month_period)
        ]
        for row in source_rows:
            if not date_in_period(row["date"], period):
                continue
            region = row_region(row, segments)
            if region not in selected_regions:
                continue
            item = dict(row)
            item["region"] = region
            rows.append(item)
    return aggregate_utilization_rows(rows)


def load_security_rows_uncached() -> list[dict[str, Any]]:
    workbook = load_workbook(DATA_FILE, data_only=True)
    if SECURITY_SHEET_NAME not in workbook.sheetnames:
        return []

    sheet = workbook[SECURITY_SHEET_NAME]
    headers = [normalize_header(cell.value) for cell in sheet[1]]
    rows: list[dict[str, Any]] = []
    for values in sheet.iter_rows(min_row=2, values_only=True):
        record = dict(zip(headers, values))
        if not any(record.values()):
            continue
        target_ip = extract_ip(first_present(record, ["目的IP", "主机IP", "受攻击IP", "资产IP", "IP"]))
        severity = normalize_header(first_present(record, ["严重程度", "告警等级", "确定性等级"]))
        attack_result = normalize_header(first_present(record, ["攻击结果", "攻击状态", "结果"]))
        rows.append(
            {
                "bureau": normalize_header(first_present(record, ["厅局名称", "单位名称", "所属单位"])),
                "system": normalize_header(first_present(record, ["系统名称", "受攻击系统", "所属业务", "所属资产组"])),
                "date": normalize_date(first_present(record, ["时间", "最近发生时间", "最早发生时间"])),
                "target_ip": target_ip,
                "attack_type": normalize_header(
                    first_present(record, ["攻击类型", "告警三级分类", "告警二级分类", "告警一级分类", "告警名称"])
                ),
                "attack_result": attack_result,
                "severity": normalize_security_severity(severity, attack_result),
            }
        )
    return rows


def load_security_rows() -> list[dict[str, Any]]:
    reset_cache_if_needed()
    if _CACHE["security_rows"] is None:
        _CACHE["security_rows"] = load_security_rows_uncached()
    return _CACHE["security_rows"]


def load_resource_rows_uncached() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    if RESOURCE_FILE.exists():
        try:
            rows.extend(load_xls_resource_rows(RESOURCE_FILE))
        except Exception as exc:
            print(f"跳过资源清单 {RESOURCE_FILE.name}：{exc}")
    if EXTRA_RESOURCE_DIR.exists():
        for path in sorted(EXTRA_RESOURCE_DIR.rglob("*")):
            if is_ignored_data_file(path) or not path.is_file():
                continue
            try:
                if path.suffix.lower() == ".xls":
                    rows.extend(load_xls_resource_rows(path))
                elif path.suffix.lower() in (".xlsx", ".xlsm"):
                    rows.extend(load_xlsx_resource_rows(path))
            except Exception as exc:
                print(f"跳过资源清单 {path.name}：{exc}")
    return rows


def resource_row_from_record(record: dict[str, Any]) -> dict[str, Any] | None:
    ip = extract_ip(first_present(record, ["ip", "IP", "主机IP"]))
    if not ip:
        return None
    return {
        "bureau": normalize_header(first_present(record, ["单位名称", "厅局名称"])),
        "system": normalize_header(first_present(record, ["系统名称"])),
        "service_type": normalize_header(first_present(record, ["服务类型"])),
        "ip": ip,
        "enabled": normalize_date(first_present(record, ["启用时间", "开始时间"])),
        "disabled": normalize_date(first_present(record, ["停用时间", "结束时间"])),
        "vcpu": first_present(record, ["vcpu", "vCPU", "CPU"]),
        "memory": first_present(record, ["内存", "内存(G)", "内存GB"]),
        "system_disk": sum_numbers(record.get("系统盘"), record.get("系统盘(G)"), record.get("高性能系统盘")),
        "data_disk": sum_numbers(record.get("数据盘"), record.get("数据盘(G)"), record.get("高性能数据盘")),
    }


def load_xls_resource_rows(path: Path) -> list[dict[str, Any]]:
    try:
        import xlrd
    except ImportError as exc:
        raise RuntimeError("缺少读取资源清单.xls 的依赖，请安装：pip install xlrd") from exc
    rows: list[dict[str, Any]] = []
    workbook = xlrd.open_workbook(str(path))
    sheet = workbook.sheet_by_index(0)
    if sheet.nrows < 2:
        return []

    headers = [normalize_header(sheet.cell_value(0, col)) for col in range(sheet.ncols)]

    def cell_value(row_index: int, name: str) -> Any:
        if name not in headers:
            return ""
        col_index = headers.index(name)
        cell = sheet.cell(row_index, col_index)
        if cell.ctype == xlrd.XL_CELL_DATE:
            return xlrd.xldate.xldate_as_datetime(cell.value, workbook.datemode)
        return cell.value

    for row_index in range(1, sheet.nrows):
        record = {name: cell_value(row_index, name) for name in headers if name}
        if not any(record.values()):
            continue
        row = resource_row_from_record(record)
        if row:
            rows.append(row)
    return rows


def load_xlsx_resource_rows(path: Path) -> list[dict[str, Any]]:
    workbook = load_workbook(path, data_only=True, read_only=True)
    rows: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        if sheet.max_row < 2:
            continue
        headers = [normalize_header(cell.value) for cell in sheet[1]]
        for values in sheet.iter_rows(min_row=2, values_only=True):
            record = dict(zip(headers, values))
            if not any(record.values()):
                continue
            row = resource_row_from_record(record)
            if row:
                rows.append(row)
    return rows


def load_resource_rows() -> list[dict[str, Any]]:
    reset_cache_if_needed()
    if _CACHE["resource_rows"] is None:
        _CACHE["resource_rows"] = load_resource_rows_uncached()
    return _CACHE["resource_rows"]


def resource_index_by_ip() -> dict[str, dict[str, str]]:
    reset_cache_if_needed()
    if _CACHE["resource_index"] is None:
        index: dict[str, dict[str, str]] = {}
        for row in load_resource_rows():
            if row["ip"] and row["ip"] not in index:
                index[row["ip"]] = {"bureau": row["bureau"], "system": row["system"]}
        _CACHE["resource_index"] = index
    return _CACHE["resource_index"]


def resource_status(row: dict[str, Any], period: dict[str, str]) -> str | None:
    start = parse_date(period["start"])
    end = parse_date(period["end"])
    enabled = parse_date(row.get("enabled", ""))
    disabled = parse_date(row.get("disabled", ""))
    if not start or not end:
        return None
    enabled_in_period = enabled is not None and start <= enabled <= end
    disabled_in_period = disabled is not None and start <= disabled <= end
    if enabled_in_period and disabled_in_period:
        return "新增并核减"
    if enabled_in_period:
        return "新增资源"
    if disabled_in_period:
        return "核减资源"
    return None


def filtered_resource_rows(
    bureau: str,
    period: dict[str, str],
    regions: list[str] | None = None,
) -> list[dict[str, Any]]:
    selected_regions = normalize_regions(regions)
    segments = load_address_segments()
    rows: list[dict[str, Any]] = []
    for row in load_resource_rows():
        if not bureau_matches(row["bureau"], bureau):
            continue
        status = resource_status(row, period)
        if not status:
            continue
        region = classify_ip_region(row["ip"], segments)
        if region not in selected_regions:
            continue
        item = dict(row)
        item["region"] = region
        item["status"] = status
        rows.append(item)
    return sorted(rows, key=lambda item: (item["system"], item["ip"], item["status"]))


def aggregate_security_rows(
    bureau: str,
    period: dict[str, str],
    regions: list[str] | None = None,
) -> list[dict[str, Any]]:
    selected_regions = normalize_regions(regions)
    segments = load_address_segments()
    grouped: OrderedDict[tuple[str, str, str, str, str], int] = OrderedDict()

    for row in load_security_rows():
        if not bureau_matches(row["bureau"], bureau) or not date_in_period(row["date"], period):
            continue
        region = classify_ip_region(row["target_ip"], segments)
        if region not in selected_regions:
            continue
        key = (
            row["target_ip"] or "-",
            region,
            row["attack_type"] or "-",
            row["system"] or "-",
            row["severity"] or "-",
        )
        grouped[key] = grouped.get(key, 0) + 1

    severity_order = {"严重": 0, "高危": 1, "高": 1, "中危": 2, "中": 2, "低危": 3, "低": 3}
    result = [
        {
            "target_ip": key[0],
            "region": key[1],
            "attack_type": key[2],
            "system": key[3],
            "severity": key[4],
            "count": count,
        }
        for key, count in grouped.items()
    ]
    return sorted(result, key=lambda item: (severity_order.get(item["severity"], 9), -item["count"], item["target_ip"]))


def group_by_system(rows: list[dict[str, Any]]) -> OrderedDict[str, list[dict[str, Any]]]:
    grouped: OrderedDict[str, list[dict[str, Any]]] = OrderedDict()
    for row in rows:
        grouped.setdefault(row["system"], []).append(row)
    return grouped


def api_options() -> dict[str, Any]:
    rows = load_rows()
    csv_rows = load_csv_utilization_rows()
    security_rows = load_security_rows()
    resource_rows = load_resource_rows()
    bureaus = sorted(
        {row["bureau"] for row in rows}
        .union({row["bureau"] for row in csv_rows if row["bureau"]})
        .union({row["bureau"] for row in security_rows if row["bureau"]})
        .union({row["bureau"] for row in resource_rows if row["bureau"]})
    )
    resource_dates = {
        value
        for row in resource_rows
        for value in (row.get("enabled", ""), row.get("disabled", ""))
        if value
    }
    dates = sorted(
        {row["date"] for row in rows}
        .union({row["date"] for row in csv_rows if row["date"]})
        .union({row["date"] for row in security_rows if row["date"]})
        .union(resource_dates)
    )
    months = sorted({value[:7] for value in dates if value})
    combinations: dict[str, list[str]] = {}
    month_combinations: dict[str, list[str]] = {}
    for bureau in bureaus:
        combinations[bureau] = sorted(
            {row["date"] for row in rows if bureau_matches(row["bureau"], bureau)}.union(
                {row["date"] for row in csv_rows if bureau_matches(row["bureau"], bureau) and row["date"]}
            ).union(
                {row["date"] for row in security_rows if bureau_matches(row["bureau"], bureau) and row["date"]}
            ).union(
                {
                    value
                    for row in resource_rows
                    if bureau_matches(row["bureau"], bureau)
                    for value in (row.get("enabled", ""), row.get("disabled", ""))
                    if value
                }
            )
        )
        month_combinations[bureau] = sorted({value[:7] for value in combinations[bureau] if value})
    return {
        "bureaus": bureaus,
        "dates": dates,
        "months": months,
        "combinations": combinations,
        "monthCombinations": month_combinations,
        "regions": REGIONS,
    }


def api_preview(bureau: str, period: dict[str, str], regions: list[str] | None = None) -> dict[str, Any]:
    rows = filtered_rows(bureau, period, regions)
    grouped = group_by_system(rows)
    security_rows = aggregate_security_rows(bureau, period, regions)
    resource_rows = filtered_resource_rows(bureau, period, regions)
    return {
        "bureau": bureau,
        "date": period["start"],
        "start": period["start"],
        "end": period["end"],
        "month": period["month"],
        "displayDate": period["label"],
        "regionLabel": region_label(regions),
        "systemCount": len(grouped),
        "rowCount": len(rows),
        "securityCount": sum(row["count"] for row in security_rows),
        "securityGroupCount": len(security_rows),
        "resourceScaleCount": len(resource_rows),
        "systems": [
            {
                "name": system,
                "rows": [
                    {
                        "ip": row["ip"],
                        "region": row["region"],
                        "cpu": metric(row["cpu"]),
                        "memory": metric(row["memory"]),
                        "disk": metric(row["disk"]),
                    }
                    for row in values
                ],
            }
            for system, values in grouped.items()
        ],
        "security": [
            {
                "targetIp": row["target_ip"],
                "region": row["region"],
                "attackType": row["attack_type"],
                "system": row["system"],
                "severity": row["severity"],
                "count": row["count"],
            }
            for row in security_rows
        ],
        "resources": [
            {
                "bureau": row["bureau"],
                "system": row["system"],
                "ip": row["ip"],
                "region": row["region"],
                "serviceType": row["service_type"],
                "vcpu": normalize_number(row["vcpu"]),
                "memory": normalize_number(row["memory"]),
                "systemDisk": normalize_number(row["system_disk"]),
                "dataDisk": normalize_number(row["data_disk"]),
                "status": row["status"],
            }
            for row in resource_rows
        ],
    }


def replace_text(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated == original:
        return

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def copy_row(table: Table, source_row: _Row) -> _Row:
    new_row = table.add_row()
    new_row._tr.getparent().remove(new_row._tr)
    table._tbl.append(copy.deepcopy(source_row._tr))
    return table.rows[-1]


def set_cell_text(cell: _Cell, text: str) -> None:
    cell.text = text


def set_table_kind(table: Table, kind: str) -> None:
    table._tbl.set("data-kind", kind)


def table_kind(table: Table) -> str:
    return table._tbl.get("data-kind", "")


def fill_table(table: Table, rows: list[dict[str, Any]]) -> None:
    if not table.rows:
        return

    if len(table.columns) < 5:
        table.add_column(Inches(1.0))
    headers = ["IP", "区域", "CPU使用率", "内存使用率", "磁盘使用率"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value)

    template_row = table.rows[1] if len(table.rows) > 1 else table.rows[0]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for row in rows:
        current = copy_row(table, template_row)
        values = [row["ip"], row["region"], metric(row["cpu"]), metric(row["memory"]), metric(row["disk"])]
        for cell, value in zip(current.cells, values):
            set_cell_text(cell, value)


def find_system_tables(document: Document) -> OrderedDict[str, Table]:
    result: OrderedDict[str, Table] = OrderedDict()
    pending_system: str | None = None
    for item in iter_block_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text and not re.search(r"[<>]", text) and "报告" not in text:
                pending_system = text
        elif isinstance(item, Table) and pending_system:
            result[pending_system] = item
            pending_system = None
    return result


def remove_block(element: Any) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(
    paragraph: Paragraph,
    size: float,
    bold: bool = False,
    color: str = COLOR_INK,
    alignment: int | None = None,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.2
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def set_paragraph_bottom_border(paragraph: Paragraph, color: str = COLOR_ACCENT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_document_background(document: Document, color: str = COLOR_PAGE_BG) -> None:
    existing = document.element.find(qn("w:background"))
    if existing is not None:
        document.element.remove(existing)
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color)
    document.element.insert(0, background)


def clear_header_waves(document: Document) -> None:
    for section in document.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for pict in list(paragraph._p.findall(".//" + qn("w:pict"))):
                parent = pict.getparent()
                if parent is not None:
                    parent.remove(pict)


def add_wave_background(document: Document) -> None:
    clear_header_waves(document)
    wave_xml = f"""
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office">
      <w:pict>
        <v:shape id="report-wave-top" o:spid="_x0000_s2049"
          style="position:absolute;left:0;text-align:left;margin-left:-52pt;margin-top:88pt;width:705pt;height:118pt;z-index:-251654144;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
          coordsize="705,118" filled="t" stroked="f"
          path="m 0,76 c 116,25 201,104 332,58 c 447,18 549,24 705,4 l 705,47 c 560,70 458,54 341,91 c 207,133 111,57 0,104 x e">
          <v:fill color="#{COLOR_WAVE}" opacity="18%"/>
        </v:shape>
        <v:shape id="report-wave-bottom" o:spid="_x0000_s2050"
          style="position:absolute;left:0;text-align:left;margin-left:22pt;margin-top:682pt;width:642pt;height:86pt;z-index:-251654143;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
          coordsize="642,86" filled="t" stroked="f"
          path="m 0,42 c 126,75 209,4 326,35 c 437,64 522,61 642,18 l 642,58 c 514,84 426,82 318,53 c 201,21 126,87 0,61 x e">
          <v:fill color="#{COLOR_WAVE_LIGHT}" opacity="30%"/>
        </v:shape>
      </w:pict>
    </w:r>
    """
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph._p.append(parse_xml(wave_xml))


def add_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run()
        set_run_font(run, size=9, color=COLOR_MUTED)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def set_cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell: _Cell, color: str = COLOR_BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        border = tc_borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            tc_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_cell_margins(cell: _Cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = tc_mar.find(qn(f"w:{name}"))
        if margin is None:
            margin = OxmlElement(f"w:{name}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_cell_width(cell: _Cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_cell(
    cell: _Cell,
    fill: str | None = None,
    text_color: str = COLOR_INK,
    bold: bool = False,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    width: int | None = None,
    font_size: float = 9.5,
    margin: tuple[int, int, int, int] = (90, 120, 90, 120),
) -> None:
    if fill:
        set_cell_shading(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell, *margin)
    if width is not None:
        set_cell_width(cell, width)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=font_size, bold=bold, color=text_color)


def style_table(table: Table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    kind = table_kind(table)
    configs = {
        "resource": {
            "widths": [1200, 1500, 1250, 720, 820, 600, 650, 700, 700, 780],
            "left": {0, 1, 2},
            "center": {3, 4, 5, 6, 7, 8, 9},
            "font": 7.0,
            "margin": (55, 55, 55, 55),
        },
        "security": {
            "widths": [1450, 850, 2200, 2700, 1000, 750],
            "left": {0, 2, 3},
            "center": {1, 4, 5},
            "font": 8.6,
            "margin": (75, 100, 75, 100),
        },
        "utilization": {
            "widths": [1800, 1050, 1350, 1350, 1350],
            "left": {0},
            "center": {1, 2, 3, 4},
            "font": 9.2,
            "margin": (80, 110, 80, 110),
        },
    }
    config = configs.get(kind, configs["utilization"])
    widths = config["widths"]
    left_columns = config["left"]
    center_columns = config["center"]

    for row_index, row in enumerate(table.rows):
        is_header = row_index == 0
        fill = COLOR_HEADER_BG if is_header else (COLOR_ROW_ALT if row_index % 2 == 0 else "FFFFFF")
        text_color = COLOR_HEADER_TEXT if is_header else COLOR_INK
        for cell_index, cell in enumerate(row.cells):
            if is_header or cell_index in center_columns:
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif cell_index in left_columns:
                align = WD_ALIGN_PARAGRAPH.LEFT
            else:
                align = WD_ALIGN_PARAGRAPH.RIGHT
            style_cell(
                cell,
                fill=fill,
                text_color=text_color,
                bold=is_header,
                align=align,
                width=widths[cell_index] if cell_index < len(widths) else None,
                font_size=config["font"],
                margin=config["margin"],
            )


def style_report(document: Document, table_map: OrderedDict[str, Table]) -> None:
    set_document_background(document)
    add_wave_background(document)
    add_page_numbers(document)
    for section in document.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.42)
        section.right_margin = Inches(0.42)

    if document.styles["Normal"]:
        normal = document.styles["Normal"]
        normal.font.name = FONT_LATIN
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        normal.font.size = Pt(10)
        normal.font.color.rgb = RGBColor.from_string(COLOR_INK)

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)

        if index == 0:
            style_paragraph(paragraph, size=22, bold=True, color=COLOR_ACCENT_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            paragraph.paragraph_format.space_after = Pt(6)
        elif index == 1:
            style_paragraph(paragraph, size=11, bold=False, color=COLOR_MUTED, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            paragraph.paragraph_format.space_after = Pt(18)
        elif text in ("云资源规模情况", "云资源使用情况", "安全态势"):
            style_paragraph(paragraph, size=13, bold=True, color=COLOR_ACCENT_DARK, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(8)
            set_paragraph_bottom_border(paragraph)
        elif text in table_map:
            style_paragraph(paragraph, size=11, bold=True, color=COLOR_ACCENT_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(6)
        else:
            style_paragraph(paragraph, size=10, color=COLOR_INK)

    for table in document.tables:
        style_table(table)


def add_security_section(document: Document, rows: list[dict[str, Any]]) -> None:
    heading = document.add_paragraph("安全态势")
    if not rows:
        paragraph = document.add_paragraph("本期所选区域未发现匹配的受攻击资产记录。")
        paragraph.paragraph_format.space_after = Pt(8)
        return

    table = document.add_table(rows=1, cols=6)
    set_table_kind(table, "security")
    headers = ["受攻击IP", "区域", "攻击类型", "受攻击系统", "严重程度", "数量"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value)

    for row in rows:
        current = table.add_row()
        values = [row["target_ip"], row["region"], row["attack_type"], row["system"], row["severity"], str(row["count"])]
        for cell, value in zip(current.cells, values):
            set_cell_text(cell, value)


def add_utilization_section(document: Document, system: str, rows: list[dict[str, Any]]) -> Table:
    document.add_paragraph(system)
    table = document.add_table(rows=1, cols=5)
    set_table_kind(table, "utilization")
    headers = ["IP", "区域", "CPU使用率", "内存使用率", "磁盘使用率"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value)
    for row in rows:
        current = table.add_row()
        values = [row["ip"], row["region"], metric(row["cpu"]), metric(row["memory"]), metric(row["disk"])]
        for cell, value in zip(current.cells, values):
            set_cell_text(cell, value)
    return table


def insert_utilization_heading(document: Document, table_map: OrderedDict[str, Table]) -> None:
    if not table_map:
        return
    first_system, first_table = next(iter(table_map.items()))
    paragraph = document.add_paragraph("云资源使用情况")
    body = document.element.body
    body.remove(paragraph._p)
    for item in document.paragraphs:
        if item.text.strip() == first_system:
            body.insert(body.index(item._p), paragraph._p)
            return
    body.insert(body.index(first_table._tbl), paragraph._p)


def add_resource_scale_section(document: Document, rows: list[dict[str, Any]]) -> None:
    document.add_paragraph("云资源规模情况")
    if not rows:
        paragraph = document.add_paragraph("本期所选区域未发现匹配的云资源规模记录。")
        paragraph.paragraph_format.space_after = Pt(8)
        return

    table = document.add_table(rows=1, cols=10)
    set_table_kind(table, "resource")
    headers = ["单位名称", "系统名称", "IP", "区域", "服务类型", "vCPU", "内存", "系统盘", "数据盘", "变化情况"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value)

    for row in rows:
        current = table.add_row()
        values = [
            row["bureau"],
            row["system"],
            row["ip"],
            row["region"],
            row["service_type"],
            normalize_number(row["vcpu"]),
            normalize_number(row["memory"]),
            normalize_number(row["system_disk"]),
            normalize_number(row["data_disk"]),
            row["status"],
        ]
        for cell, value in zip(current.cells, values):
            set_cell_text(cell, value)


def generate_report(bureau: str, period: dict[str, str], regions: list[str] | None = None) -> tuple[bytes, str]:
    rows = filtered_rows(bureau, period, regions)
    grouped = group_by_system(rows)
    security_rows = aggregate_security_rows(bureau, period, regions)
    resource_rows = filtered_resource_rows(bureau, period, regions)
    if not rows and not security_rows and not resource_rows:
        raise ValueError("没有找到匹配的报表数据。")
    document = Document(TEMPLATE_FILE)
    replacements = {
        "<厅局名称>": bureau,
        "<时间>": f"{period['label']}（{region_label(regions)}）",
    }

    for paragraph in document.paragraphs:
        replace_text(paragraph, replacements)

    table_map = find_system_tables(document)
    for system, table in table_map.items():
        system_rows = grouped.get(system, [])
        if system_rows:
            fill_table(table, system_rows)
            set_table_kind(table, "utilization")
        else:
            remove_block(table._tbl)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in table_map and text not in grouped:
            remove_block(paragraph._p)

    visible_table_map = OrderedDict((system, table) for system, table in table_map.items() if system in grouped)
    for system, system_rows in grouped.items():
        if system not in visible_table_map:
            visible_table_map[system] = add_utilization_section(document, system, system_rows)
    insert_utilization_heading(document, visible_table_map)
    add_resource_scale_section(document, resource_rows)
    add_security_section(document, security_rows)
    style_report(document, visible_table_map)

    output = io.BytesIO()
    document.save(output)
    filename = f"{bureau}-{period['filename']}-{region_label(regions)}运行报告.docx"
    return output.getvalue(), filename


INDEX_HTML = """<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>报表生成</title>
  <style>
    :root {
      color-scheme: light;
      --bg: #f6f7f9;
      --panel: #ffffff;
      --ink: #20242a;
      --muted: #667085;
      --line: #d9dee7;
      --accent: #1f6f68;
      --accent-2: #b45309;
      --danger: #b42318;
      --shadow: 0 16px 40px rgba(15, 23, 42, 0.08);
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      min-height: 100vh;
      background: var(--bg);
      color: var(--ink);
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
    }
    header {
      background: #24313d;
      color: #fff;
      padding: 22px clamp(18px, 4vw, 46px);
      border-bottom: 4px solid var(--accent);
    }
    h1 {
      margin: 0;
      font-size: clamp(24px, 3vw, 34px);
      font-weight: 750;
      letter-spacing: 0;
    }
    main {
      width: min(1180px, calc(100% - 32px));
      margin: 26px auto 48px;
      display: grid;
      gap: 18px;
    }
    .toolbar {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: var(--shadow);
      padding: 18px;
      display: grid;
      grid-template-columns: minmax(220px, 1fr) minmax(360px, 1.25fr) minmax(220px, .9fr) auto;
      align-items: end;
      gap: 14px;
    }
    label {
      display: grid;
      gap: 7px;
      color: var(--muted);
      font-size: 13px;
      font-weight: 650;
    }
    .bureau-field {
      position: relative;
      display: grid;
      gap: 7px;
      color: var(--muted);
      font-size: 13px;
      font-weight: 650;
    }
    .bureau-combobox { position: relative; }
    .bureau-options {
      position: absolute;
      z-index: 20;
      top: calc(100% + 4px);
      left: 0;
      right: 0;
      max-height: 260px;
      overflow-y: auto;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: var(--shadow);
      padding: 6px;
    }
    .bureau-option {
      width: 100%;
      min-height: 34px;
      display: block;
      text-align: left;
      border: 0;
      border-radius: 6px;
      background: #fff;
      color: var(--ink);
      padding: 7px 9px;
      font-size: 14px;
      font-weight: 600;
      white-space: normal;
    }
    .bureau-option:hover, .bureau-option.active {
      background: #eef4f3;
      color: var(--accent);
    }
    .region-field {
      display: grid;
      gap: 7px;
      color: var(--muted);
      font-size: 13px;
      font-weight: 650;
    }
    .region-options {
      min-height: 42px;
      display: flex;
      align-items: center;
      gap: 8px;
      flex-wrap: wrap;
    }
    .region-options label {
      display: inline-flex;
      align-items: center;
      gap: 6px;
      min-height: 34px;
      padding: 0 10px;
      border: 1px solid var(--line);
      border-radius: 6px;
      background: #fff;
      color: var(--ink);
      font-size: 14px;
      font-weight: 650;
      cursor: pointer;
    }
    .region-options input {
      width: 16px;
      height: 16px;
      accent-color: var(--accent);
    }
    select, input, button {
      min-height: 42px;
      border-radius: 6px;
      font: inherit;
    }
    select, input {
      width: 100%;
      border: 1px solid var(--line);
      background: #fff;
      color: var(--ink);
      padding: 0 12px;
    }
    .time-field {
      display: grid;
      gap: 8px;
      color: var(--muted);
      font-size: 13px;
      font-weight: 650;
    }
    .time-controls {
      display: grid;
      grid-template-columns: 148px minmax(150px, 190px);
      gap: 8px;
      align-items: center;
      justify-content: start;
    }
    .segmented {
      display: grid;
      grid-template-columns: 1fr 1fr;
      min-height: 42px;
      border: 1px solid var(--line);
      border-radius: 6px;
      overflow: hidden;
      background: #fff;
    }
    .segmented button {
      min-height: 40px;
      border-radius: 0;
      background: #fff;
      color: var(--ink);
      padding: 0 10px;
      border-right: 1px solid var(--line);
      font-weight: 700;
    }
    .segmented button:last-child { border-right: 0; }
    .segmented button.active {
      background: var(--accent);
      color: #fff;
    }
    .range-inputs {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 8px;
      grid-column: 2 / span 1;
      width: min(100%, 300px);
    }
    #month {
      width: 180px;
      min-width: 180px;
    }
    .hidden { display: none; }
    .actions {
      display: flex;
      align-items: center;
      gap: 8px;
      min-height: 42px;
    }
    button {
      border: 0;
      background: var(--accent);
      color: #fff;
      padding: 0 18px;
      font-weight: 700;
      cursor: pointer;
      white-space: nowrap;
    }
    button:disabled {
      opacity: .55;
      cursor: wait;
    }
    .secondary {
      background: #eef4f3;
      color: var(--accent);
      border: 1px solid #b9d5d1;
    }
    .status {
      min-height: 24px;
      color: var(--muted);
      font-size: 14px;
    }
    .status.error { color: var(--danger); }
    .summary {
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 12px;
    }
    .stat {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 15px 16px;
    }
    .stat span {
      display: block;
      color: var(--muted);
      font-size: 13px;
      margin-bottom: 5px;
    }
    .stat strong {
      font-size: 24px;
      color: #28313b;
    }
    .systems {
      display: grid;
      gap: 16px;
    }
    .system {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: hidden;
    }
    .system h2 {
      margin: 0;
      padding: 14px 16px;
      font-size: 17px;
      background: #eef4f3;
      border-bottom: 1px solid var(--line);
      letter-spacing: 0;
    }
    .table-wrap { overflow-x: auto; }
    table {
      width: 100%;
      border-collapse: collapse;
      min-width: 620px;
    }
    th, td {
      padding: 10px 12px;
      border-bottom: 1px solid #e7ebf0;
      text-align: left;
      font-size: 14px;
      white-space: nowrap;
    }
    th {
      background: #fbfcfd;
      color: #4b5563;
      font-weight: 700;
    }
    td:nth-child(n+2) { color: var(--accent-2); font-variant-numeric: tabular-nums; }
    @media (max-width: 760px) {
      .toolbar, .summary { grid-template-columns: 1fr; }
      button { width: 100%; }
      header { padding: 18px 16px; }
    }
  </style>
</head>
<body>
  <header>
    <h1>月度运行报告生成器</h1>
  </header>
  <main>
    <section class="toolbar" aria-label="报表条件">
      <div class="bureau-field">厅局
        <div class="bureau-combobox">
          <input id="bureau-search" type="search" autocomplete="off" placeholder="输入厅局名称搜索" />
          <div id="bureau-options" class="bureau-options hidden"></div>
        </div>
      </div>
      <div class="time-field">时间
        <div class="time-controls">
          <div class="segmented" aria-label="时间模式">
            <button id="month-mode" class="active" type="button">按月</button>
            <button id="range-mode" type="button">范围</button>
          </div>
          <input id="month" type="month" />
          <div id="range-inputs" class="range-inputs hidden">
            <input id="start-date" type="date" />
            <input id="end-date" type="date" />
          </div>
        </div>
      </div>
      <div class="region-field">区域
        <div class="region-options" id="regions"></div>
      </div>
      <div class="actions">
        <button id="query" class="secondary" type="button">查询</button>
        <button id="generate" type="button" disabled>生成报表</button>
      </div>
    </section>
    <div id="status" class="status"></div>
    <section class="summary" aria-label="数据概览">
      <div class="stat"><span>厅局名称</span><strong id="stat-bureau">-</strong></div>
      <div class="stat"><span>报表时间</span><strong id="stat-date">-</strong></div>
      <div class="stat"><span>资源规模 / 安全态势</span><strong id="stat-count">-</strong></div>
    </section>
    <section id="systems" class="systems" aria-label="预览数据"></section>
  </main>
  <script>
    const bureauSearchEl = document.querySelector("#bureau-search");
    const bureauOptionsEl = document.querySelector("#bureau-options");
    const monthEl = document.querySelector("#month");
    const startDateEl = document.querySelector("#start-date");
    const endDateEl = document.querySelector("#end-date");
    const monthModeEl = document.querySelector("#month-mode");
    const rangeModeEl = document.querySelector("#range-mode");
    const rangeInputsEl = document.querySelector("#range-inputs");
    const queryEl = document.querySelector("#query");
    const buttonEl = document.querySelector("#generate");
    const regionsEl = document.querySelector("#regions");
    const statusEl = document.querySelector("#status");
    const systemsEl = document.querySelector("#systems");
    const statBureau = document.querySelector("#stat-bureau");
    const statDate = document.querySelector("#stat-date");
    const statCount = document.querySelector("#stat-count");
    let combinations = {};
    let monthCombinations = {};
    let bureaus = [];
    let selectedBureau = "";
    let regions = [];
    let timeMode = "month";
    let lastQueryPayload = null;

    function setStatus(message, isError = false) {
      statusEl.textContent = message;
      statusEl.classList.toggle("error", isError);
    }

    function escapeHtml(value) {
      return String(value).replace(/[&<>"']/g, char => ({
        "&": "&amp;",
        "<": "&lt;",
        ">": "&gt;",
        '"': "&quot;",
        "'": "&#39;"
      }[char]));
    }

    function filteredBureaus(keyword) {
      const text = keyword.trim().toLowerCase();
      if (!text) return bureaus.slice(0, 30);
      return bureaus.filter(name => name.toLowerCase().includes(text)).slice(0, 30);
    }

    function renderBureauOptions(keyword = bureauSearchEl.value) {
      const matches = filteredBureaus(keyword);
      bureauOptionsEl.innerHTML = matches.length ? matches.map((name, index) => `
        <button class="bureau-option${index === 0 ? " active" : ""}" type="button" data-name="${escapeHtml(name)}">${escapeHtml(name)}</button>
      `).join("") : `<button class="bureau-option" type="button" disabled>无匹配厅局</button>`;
      bureauOptionsEl.classList.remove("hidden");
      bureauOptionsEl.querySelectorAll("button[data-name]").forEach(button => {
        button.addEventListener("click", () => selectBureau(button.dataset.name));
      });
    }

    function selectBureau(name) {
      selectedBureau = name;
      bureauSearchEl.value = name;
      bureauOptionsEl.classList.add("hidden");
      updateDateOptions();
      markDirty("条件已变化，请点击查询。");
    }

    function fillRegions(values) {
      regions = values.length ? values : ["信创", "非信创"];
      regionsEl.innerHTML = regions.map(value => `
        <label>
          <input type="checkbox" name="region" value="${escapeHtml(value)}" checked />
          ${escapeHtml(value)}
        </label>
      `).join("");
      regionsEl.querySelectorAll("input").forEach(input => {
        input.addEventListener("change", () => markDirty());
      });
    }

    function selectedRegions() {
      const values = Array.from(regionsEl.querySelectorAll("input:checked")).map(input => input.value);
      if (!values.length) {
        regionsEl.querySelectorAll("input").forEach(input => input.checked = true);
        return regions.slice();
      }
      return values;
    }

    async function loadOptions() {
      const response = await fetch("/api/options");
      if (!response.ok) throw new Error("读取选项失败");
      const data = await response.json();
      combinations = data.combinations || {};
      monthCombinations = data.monthCombinations || {};
      bureaus = data.bureaus || [];
      if (bureaus.length) selectBureau(bureaus[0]);
      fillRegions(data.regions || []);
      markDirty("请选择条件后点击查询。");
    }

    function updateDateOptions() {
      const bureau = selectedBureau;
      const dates = combinations[bureau] || [];
      const months = monthCombinations[bureau] || [];
      monthEl.value = months[months.length - 1] || "";
      startDateEl.value = dates[0] || "";
      endDateEl.value = dates[dates.length - 1] || dates[0] || "";
    }

    function setTimeMode(mode) {
      timeMode = mode;
      monthModeEl.classList.toggle("active", mode === "month");
      rangeModeEl.classList.toggle("active", mode === "range");
      monthEl.classList.toggle("hidden", mode !== "month");
      rangeInputsEl.classList.toggle("hidden", mode !== "range");
      markDirty();
    }

    function selectedPeriod() {
      return {
        mode: timeMode,
        month: monthEl.value,
        start: startDateEl.value,
        end: endDateEl.value
      };
    }

    function currentPayload() {
      return { bureau: selectedBureau, ...selectedPeriod(), regions: selectedRegions() };
    }

    function markDirty(message = "条件已变化，请点击查询。") {
      lastQueryPayload = null;
      buttonEl.disabled = true;
      statBureau.textContent = "-";
      statDate.textContent = "-";
      statCount.textContent = "-";
      systemsEl.innerHTML = "";
      setStatus(message);
    }

    async function loadPreview() {
      const bureau = selectedBureau;
      const period = selectedPeriod();
      const regionValues = selectedRegions();
      if (!bureau || (timeMode === "month" && !period.month) || (timeMode === "range" && (!period.start || !period.end))) {
        setStatus("没有可用的报表数据。", true);
        return;
      }
      setStatus("正在读取预览数据...");
      const params = new URLSearchParams({ bureau, mode: period.mode });
      if (period.mode === "month") {
        params.set("month", period.month);
      } else {
        params.set("start", period.start);
        params.set("end", period.end);
      }
      regionValues.forEach(region => params.append("region", region));
      const response = await fetch(`/api/preview?${params.toString()}`);
      if (!response.ok) throw new Error("读取预览失败");
      const data = await response.json();
      lastQueryPayload = currentPayload();
      buttonEl.disabled = false;
      statBureau.textContent = data.bureau || "-";
      statDate.textContent = data.displayDate || data.date || "-";
      statCount.textContent = `${data.resourceScaleCount} / ${data.securityCount}`;
      systemsEl.innerHTML = data.systems.map(system => `
        <article class="system">
          <h2>${escapeHtml(system.name)}</h2>
          <div class="table-wrap">
            <table>
              <thead><tr><th>IP</th><th>区域</th><th>CPU使用率</th><th>内存使用率</th><th>磁盘使用率</th></tr></thead>
              <tbody>
                ${system.rows.map(row => `<tr><td>${escapeHtml(row.ip)}</td><td>${escapeHtml(row.region)}</td><td>${escapeHtml(row.cpu)}</td><td>${escapeHtml(row.memory)}</td><td>${escapeHtml(row.disk)}</td></tr>`).join("")}
              </tbody>
            </table>
          </div>
        </article>
      `).join("");
      const resourceHtml = `
        <article class="system">
          <h2>云资源规模情况</h2>
          <div class="table-wrap">
            <table>
              <thead><tr><th>单位名称</th><th>系统名称</th><th>IP</th><th>区域</th><th>服务类型</th><th>vCPU</th><th>内存</th><th>系统盘</th><th>数据盘</th><th>变化情况</th></tr></thead>
              <tbody>
                ${(data.resources || []).length ? data.resources.map(row => `<tr><td>${escapeHtml(row.bureau)}</td><td>${escapeHtml(row.system)}</td><td>${escapeHtml(row.ip)}</td><td>${escapeHtml(row.region)}</td><td>${escapeHtml(row.serviceType)}</td><td>${escapeHtml(row.vcpu)}</td><td>${escapeHtml(row.memory)}</td><td>${escapeHtml(row.systemDisk)}</td><td>${escapeHtml(row.dataDisk)}</td><td>${escapeHtml(row.status)}</td></tr>`).join("") : `<tr><td colspan="10">所选条件下暂无云资源规模变化数据</td></tr>`}
              </tbody>
            </table>
          </div>
        </article>
      `;
      systemsEl.insertAdjacentHTML("beforeend", resourceHtml);
      const securityHtml = `
        <article class="system">
          <h2>安全态势</h2>
          <div class="table-wrap">
            <table>
              <thead><tr><th>受攻击IP</th><th>区域</th><th>攻击类型</th><th>受攻击系统</th><th>严重程度</th><th>数量</th></tr></thead>
              <tbody>
                ${(data.security || []).length ? data.security.map(row => `<tr><td>${escapeHtml(row.targetIp)}</td><td>${escapeHtml(row.region)}</td><td>${escapeHtml(row.attackType)}</td><td>${escapeHtml(row.system)}</td><td>${escapeHtml(row.severity)}</td><td>${escapeHtml(row.count)}</td></tr>`).join("") : `<tr><td colspan="6">所选条件下暂无安全态势数据</td></tr>`}
              </tbody>
            </table>
          </div>
        </article>
      `;
      systemsEl.insertAdjacentHTML("beforeend", securityHtml);
      setStatus(`已载入 ${data.rowCount} 条利用率数据，云资源 ${data.resourceScaleCount} 条，安全态势 ${data.securityCount} 条。`);
    }

    async function generateReport() {
      if (!lastQueryPayload) {
        setStatus("请先点击查询，再生成报表。", true);
        return;
      }
      buttonEl.disabled = true;
      setStatus("正在生成 Word 报表...");
      try {
        const response = await fetch("/generate", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify(lastQueryPayload)
        });
        if (!response.ok) {
          const payload = await response.json().catch(() => ({}));
          throw new Error(payload.error || "生成失败");
        }
        const blob = await response.blob();
        const disposition = response.headers.get("Content-Disposition") || "";
        const matched = disposition.match(/filename\\*=UTF-8''([^;]+)/);
        const filename = matched ? decodeURIComponent(matched[1]) : "运行报告.docx";
        const url = URL.createObjectURL(blob);
        const link = document.createElement("a");
        link.href = url;
        link.download = filename;
        document.body.appendChild(link);
        link.click();
        link.remove();
        URL.revokeObjectURL(url);
        setStatus("报表已生成并开始下载。");
      } catch (error) {
        setStatus(error.message, true);
      } finally {
        buttonEl.disabled = !lastQueryPayload;
      }
    }

    bureauSearchEl.addEventListener("input", () => {
      selectedBureau = bureaus.includes(bureauSearchEl.value) ? bureauSearchEl.value : "";
      renderBureauOptions();
      markDirty("请选择匹配的厅局后点击查询。");
    });
    bureauSearchEl.addEventListener("focus", () => renderBureauOptions());
    bureauSearchEl.addEventListener("keydown", event => {
      if (event.key === "Enter") {
        const first = bureauOptionsEl.querySelector("button[data-name]");
        if (first) selectBureau(first.dataset.name);
        event.preventDefault();
      } else if (event.key === "Escape") {
        bureauOptionsEl.classList.add("hidden");
      }
    });
    document.addEventListener("click", event => {
      if (!event.target.closest(".bureau-combobox")) bureauOptionsEl.classList.add("hidden");
    });
    monthEl.addEventListener("change", () => markDirty());
    startDateEl.addEventListener("change", () => markDirty());
    endDateEl.addEventListener("change", () => markDirty());
    monthModeEl.addEventListener("click", () => setTimeMode("month"));
    rangeModeEl.addEventListener("click", () => setTimeMode("range"));
    queryEl.addEventListener("click", () => loadPreview().catch(error => setStatus(error.message, true)));
    buttonEl.addEventListener("click", generateReport);
    loadOptions().catch(error => setStatus(error.message, true));
  </script>
</body>
</html>
"""


def response(
    start_response,
    status: str,
    body: bytes,
    content_type: str,
    headers: list[tuple[str, str]] | None = None,
) -> list[bytes]:
    final_headers = [
        ("Content-Type", content_type),
        ("Content-Length", str(len(body))),
    ]
    if headers:
        final_headers.extend(headers)
    start_response(status, final_headers)
    return [body]


def json_response(start_response, payload: dict[str, Any], status: str = "200 OK") -> list[bytes]:
    return response(
        start_response,
        status,
        json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        "application/json; charset=utf-8",
    )


def app(environ, start_response):
    try:
        method = environ["REQUEST_METHOD"]
        parsed = urlparse(environ.get("PATH_INFO", "/"))
        path = parsed.path

        if method == "GET" and path == "/":
            return response(start_response, "200 OK", INDEX_HTML.encode("utf-8"), "text/html; charset=utf-8")

        if method == "GET" and path == "/api/options":
            return json_response(start_response, api_options())

        if method == "GET" and path == "/api/preview":
            query = parse_qs(environ.get("QUERY_STRING", ""))
            bureau = query.get("bureau", [""])[0]
            period = normalize_period(
                mode=query.get("mode", ["month"])[0],
                month=query.get("month", [""])[0],
                start=query.get("start", [""])[0],
                end=query.get("end", [""])[0],
                date_value=query.get("date", [""])[0],
            )
            regions = query.get("region", [])
            return json_response(start_response, api_preview(bureau, period, regions))

        if method == "POST" and path == "/generate":
            length = int(environ.get("CONTENT_LENGTH") or 0)
            payload = json.loads(environ["wsgi.input"].read(length) or b"{}")
            period = normalize_period(
                mode=payload.get("mode") or "month",
                month=payload.get("month") or "",
                start=payload.get("start") or "",
                end=payload.get("end") or "",
                date_value=payload.get("date") or "",
            )
            data, filename = generate_report(
                str(payload.get("bureau", "")),
                period,
                payload.get("regions") or REGIONS,
            )
            return response(
                start_response,
                "200 OK",
                data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                [("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")],
            )

        return json_response(start_response, {"error": "接口不存在。"}, "404 Not Found")
    except Exception as exc:
        return json_response(start_response, {"error": str(exc)}, "500 Internal Server Error")


class QuietRequestHandler(WSGIRequestHandler):
    def log_message(self, format, *args):
        return


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "8000"))
    host = os.environ.get("HOST", "0.0.0.0")
    print(f"报表生成服务已启动：http://localhost:{port}")
    with make_server(host, port, app, handler_class=QuietRequestHandler) as server:
        server.serve_forever()
